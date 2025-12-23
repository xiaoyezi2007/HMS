from datetime import date, datetime, timedelta

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import func, or_
from sqlalchemy.exc import IntegrityError
from sqlmodel import select, SQLModel
from typing import List, Optional
import io
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import get_session
# --- 修改点：从 deps 导入，不再依赖 patient_service ---
from app.api.deps import get_current_user_phone
from app.models.user import UserAccount, UserRole
from app.models.hospital import (
    Doctor,
    Registration,
    RegStatus,
    MedicalRecord,
    Patient,
    Examination,
    Payment,
    PaymentType,
    Prescription,
    PrescriptionDetail,
    Ward,
    Hospitalization,
    ExamResult,
    Medicine,
    NurseTask,
)
from app.schemas.hospital import MedicalRecordCreate
from app.schemas.hospital import ExaminationCreate, NurseTaskBatchCreate, NurseTaskPlan
import random

router = APIRouter()


class HospitalizePayload(SQLModel):
    ward_id: int
    hosp_doctor_id: Optional[int] = None


# --- 辅助函数：获取当前登录的医生对象 ---
async def get_current_doctor(
        phone: str = Depends(get_current_user_phone),
        session: AsyncSession = Depends(get_session)
) -> Doctor:
    user_stmt = select(UserAccount).where(UserAccount.phone == phone)
    user = (await session.execute(user_stmt)).scalars().first()
    if not user or user.role != UserRole.DOCTOR:
        raise HTTPException(status_code=403, detail="无权访问：仅限医生操作")

    doc_stmt = select(Doctor).where(Doctor.phone == phone)
    doctor = (await session.execute(doc_stmt)).scalars().first()
    if not doctor:
        raise HTTPException(status_code=404, detail="未找到您的医生档案信息")

    return doctor


def format_datetime(dt):
    if not dt:
        return ""
    return dt.strftime("%Y-%m-%d %H:%M")


def build_task_detail(plan: NurseTaskPlan) -> Optional[str]:
    detail_parts: List[str] = []
    if plan.detail and plan.detail.strip():
        detail_parts.append(plan.detail.strip())

    if plan.interval_days:
        detail_parts.append(f"频次：每 {plan.interval_days} 天 1 次，持续 {plan.duration_days} 天")
    else:
        times = plan.times_per_day or 1
        detail_parts.append(f"频次：每天 {times} 次，持续 {plan.duration_days} 天")

    return "；".join(detail_parts) if detail_parts else None


def expand_task_schedule(plan: NurseTaskPlan) -> List[datetime]:
    occurrences: List[datetime] = []
    if plan.interval_days:
        interval = plan.interval_days
        current_time = plan.start_time
        days_elapsed = 0
        while days_elapsed < plan.duration_days:
            occurrences.append(current_time)
            days_elapsed += interval
            current_time += timedelta(days=interval)
    else:
        times_per_day = plan.times_per_day or 1
        start = plan.start_time
        target_count = plan.duration_days * times_per_day
        day_offset = 0
        while len(occurrences) < target_count:
            day_base = (start + timedelta(days=day_offset)).replace(hour=0, minute=0, second=0, microsecond=0)

            if times_per_day == 1:
                dt = (start if day_offset == 0 else start + timedelta(days=day_offset))
                if dt >= start:
                    occurrences.append(dt)
            else:
                daily_times = [13, 20] if times_per_day == 2 else [8, 13, 19]
                for t in daily_times:
                    dt = day_base + timedelta(hours=t)
                    if dt >= start:
                        occurrences.append(dt)

            day_offset += 1

    return occurrences


SERVICE_FEE_RANGE = {
    "针灸": (80, 200),
    "手术": (800, 1500),
}


def estimate_service_fee(task_type: str) -> Optional[float]:
    fee_range = SERVICE_FEE_RANGE.get(task_type)
    if not fee_range:
        return None
    return round(random.uniform(*fee_range), 2)


# --- 接口 1: 医生排班 ---
@router.get("/schedule", response_model=List[Registration])
async def get_my_schedule(
        doctor: Doctor = Depends(get_current_doctor),
        session: AsyncSession = Depends(get_session)
):
    # Auto-expire overdue registrations (WAITING and visit_date has passed)
    try:
        today = date.today()
        overdue_stmt = select(Registration).where(
            Registration.doctor_id == doctor.doctor_id,
            Registration.status == RegStatus.WAITING,
            Registration.visit_date < today,
        )
        overdue_regs = (await session.execute(overdue_stmt)).scalars().all()
        for reg in overdue_regs:
            reg.status = RegStatus.EXPIRED
            session.add(reg)

            pay_stmt = (
                select(Payment)
                .where(Payment.patient_id == reg.patient_id)
                .where(
                    or_(
                        Payment.type == PaymentType.REGISTRATION,
                        Payment.type == "挂号费",
                        Payment.type == "REGISTRATION",
                    )
                )
                .where(Payment.reg_id == reg.reg_id)
                .order_by(Payment.time.desc(), Payment.payment_id.desc())
            )
            pay = (await session.execute(pay_stmt)).scalars().first()
            if pay and getattr(pay, "status", "未缴费") != "已缴费":
                pay.status = "已取消"
                session.add(pay)

        if overdue_regs:
            await session.commit()
    except Exception:
        # best-effort; schedule should still load
        pass

    # 返回排队中和办理中的挂号，医生在完成办理后挂号会被标记为已结束并从列表中消失
    stmt = select(Registration).where(
        Registration.doctor_id == doctor.doctor_id,
        Registration.status.in_([RegStatus.WAITING, RegStatus.IN_PROGRESS])
    )
    result = await session.execute(stmt)
    return result.scalars().all()


# --- 接口 2: 接诊 ---
@router.post("/consultations/{reg_id}", response_model=MedicalRecord)
async def create_medical_record(
        reg_id: int,
        record_in: MedicalRecordCreate,
        doctor: Doctor = Depends(get_current_doctor),
        session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()

    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能处理自己的病人")
    # 必须先开始办理（状态为办理中），才能书写病历
    if registration.status != RegStatus.IN_PROGRESS:
        raise HTTPException(status_code=400, detail="请先开始办理，再书写病历")

    # 如果已经存在病历，则将其视为更新（修改病历），保持 1:1 关系
    existing = (await session.execute(select(MedicalRecord).where(MedicalRecord.reg_id == reg_id))).scalars().first()
    if existing:
        existing.complaint = record_in.complaint
        existing.diagnosis = record_in.diagnosis
        existing.suggestion = record_in.suggestion
        session.add(existing)
        await session.commit()
        await session.refresh(existing)
        return existing

    new_record = MedicalRecord(
        reg_id=reg_id,
        **record_in.dict()
    )
    session.add(new_record)

    # 不在此处修改挂号状态，挂号的结束由医生在完成办理时统一设置为已完成
    await session.commit()
    await session.refresh(new_record)
    return new_record

# --- 新：为病历添加检查记录（医生操作，会随机生成结果） ---
@router.post("/consultations/{reg_id}/exams", response_model=Examination)
async def create_examination(
    reg_id: int,
    exam_in: ExaminationCreate,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    # 验证挂号与医生关联
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能为自己的病人添加检查")

    # 必须存在病历才能添加检查
    rec_stmt = select(MedicalRecord).where(MedicalRecord.reg_id == reg_id)
    record = (await session.execute(rec_stmt)).scalars().first()
    if not record:
        raise HTTPException(status_code=404, detail="请先为该挂号建立病历，再添加检查")

    # 随机生成检查结果
    choices = [r.value for r in ExamResult]
    result = random.choice(choices)

    new_exam = Examination(
        type=exam_in.type,
        result=result,
        record_id=record.record_id
    )
    session.add(new_exam)
    await session.commit()
    await session.refresh(new_exam)

    return new_exam


# --- 新：列出挂号对应的检查记录（医生操作） ---
@router.get("/consultations/{reg_id}/exams")
async def list_examinations(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    # 验证挂号与医生关联
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能查看自己的病人的检查记录")

    rec_stmt = select(MedicalRecord).where(MedicalRecord.reg_id == reg_id)
    record = (await session.execute(rec_stmt)).scalars().first()
    if not record:
        return []

    exam_stmt = select(Examination).where(Examination.record_id == record.record_id)
    exams = (await session.execute(exam_stmt)).scalars().all()
    return exams


@router.get("/wards")
async def list_department_wards(
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Ward).where(Ward.dept_id == doctor.dept_id)
    wards = (await session.execute(stmt)).scalars().all()
    out = []
    for ward in wards:
        count_stmt = select(func.count()).where(
            Hospitalization.ward_id == ward.ward_id,
            Hospitalization.status == "在院"
        )
        occupied = (await session.execute(count_stmt)).scalar_one() or 0
        available = max(ward.bed_count - occupied, 0)
        out.append({
            "ward_id": ward.ward_id,
            "type": ward.type,
            "bed_count": ward.bed_count,
            "occupied": occupied,
            "available": available,
            "is_full": occupied >= ward.bed_count
        })
    return out


@router.get("/inpatients")
async def list_my_inpatients(
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = (
        select(Hospitalization, Ward, MedicalRecord, Registration, Patient)
        .join(Ward, Hospitalization.ward_id == Ward.ward_id)
        .join(MedicalRecord, Hospitalization.record_id == MedicalRecord.record_id)
        .join(Registration, MedicalRecord.reg_id == Registration.reg_id)
        .join(Patient, Registration.patient_id == Patient.patient_id)
        .where(Hospitalization.status == "在院")
        .where(Hospitalization.hosp_doctor_id == doctor.doctor_id)
    )

    rows = await session.execute(stmt)
    now = datetime.now()
    payload = []
    for hosp, ward, _, reg, patient in rows.all():
        stay_hours = max((now - hosp.in_date).total_seconds() / 3600, 0.0)
        payload.append({
            "hosp_id": hosp.hosp_id,
            "patient_id": patient.patient_id,
            "patient_name": patient.name,
            "ward_id": ward.ward_id,
            "ward_type": ward.type,
            "in_date": hosp.in_date,
            "stay_hours": stay_hours,
            "reg_id": reg.reg_id,
        })
    return payload


@router.get("/patients/{patient_id}/registrations/history")
async def get_patient_registration_history(
    patient_id: int,
    range: str = "current",
    current_reg_id: Optional[int] = None,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    patient = await session.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="患者不存在")

    allowed_ranges = {"current", "7d", "30d"}
    if range not in allowed_ranges:
        raise HTTPException(status_code=400, detail="range 仅支持 current、7d、30d")

    stmt = (
        select(Registration, MedicalRecord)
        .outerjoin(MedicalRecord, MedicalRecord.reg_id == Registration.reg_id)
        .where(Registration.patient_id == patient_id)
        .where(Registration.doctor_id == doctor.doctor_id)
    )

    if range == "current":
        if not current_reg_id:
            raise HTTPException(status_code=400, detail="当前范围需要指定当前挂号")
        stmt = stmt.where(Registration.reg_id == current_reg_id)
    else:
        days = 7 if range == "7d" else 30
        since = datetime.now() - timedelta(days=days)
        stmt = stmt.where(Registration.status == RegStatus.FINISHED)
        stmt = stmt.where(Registration.reg_date >= since)

    stmt = stmt.order_by(Registration.reg_date.desc())
    rows = await session.execute(stmt)
    entries = []
    raw_rows = rows.all()
    for reg, record in raw_rows:
        status_value = reg.status.value if hasattr(reg.status, "value") else reg.status
        entries.append({
            "reg_id": reg.reg_id,
            "reg_date": reg.reg_date,
            "visit_date": reg.visit_date,
            "status": status_value,
            "reg_type": reg.reg_type,
            "fee": float(reg.fee or 0),
            "patient_id": reg.patient_id,
            "doctor_id": reg.doctor_id,
            "is_current": bool(current_reg_id and reg.reg_id == current_reg_id),
            "record": (
                {
                    "record_id": record.record_id,
                    "complaint": record.complaint,
                    "diagnosis": record.diagnosis,
                    "suggestion": record.suggestion
                }
                if record else None
            )
        })

    if range == "current" and not entries:
        raise HTTPException(status_code=404, detail="未找到当前挂号记录")

    return entries


@router.get("/registrations/{reg_id}/detail")
async def get_registration_detail_for_doctor(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    registration = await session.get(Registration, reg_id)
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能查看自己的挂号记录")

    record = (
        await session.execute(select(MedicalRecord).where(MedicalRecord.reg_id == reg_id))
    ).scalars().first()

    prescriptions_payload = []
    if record:
        pres_stmt = select(Prescription).where(Prescription.record_id == record.record_id)
        prescriptions = (await session.execute(pres_stmt)).scalars().all()
        for pres in prescriptions:
            details_stmt = select(PrescriptionDetail).where(PrescriptionDetail.pres_id == pres.pres_id)
            details = (await session.execute(details_stmt)).scalars().all()
            enriched = []
            for detail in details:
                medicine = await session.get(Medicine, detail.medicine_id)
                enriched.append({
                    "detail_id": detail.detail_id,
                    "medicine_id": detail.medicine_id,
                    "medicine_name": medicine.name if medicine else None,
                    "quantity": detail.quantity,
                    "usage": detail.usage,
                })
            prescriptions_payload.append({
                "pres_id": pres.pres_id,
                "create_time": pres.create_time,
                "total_amount": pres.total_amount,
                "details": enriched,
            })

    exams_payload = []
    if record:
        exam_stmt = select(Examination).where(Examination.record_id == record.record_id)
        exams = (await session.execute(exam_stmt)).scalars().all()
        for exam in exams:
            exams_payload.append({
                "exam_id": exam.exam_id,
                "type": exam.type,
                "result": exam.result,
                "date": exam.date,
                "record_id": exam.record_id,
                "reg_id": record.reg_id,
            })

    return {
        "registration": registration,
        "record": record,
        "prescriptions": prescriptions_payload,
        "exams": exams_payload,
        "admissions": []
    }


@router.post("/hospitalizations/{hosp_id}/tasks")
async def create_nurse_tasks(
    hosp_id: int,
    payload: NurseTaskBatchCreate,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    if not payload.plans:
        raise HTTPException(status_code=400, detail="请至少添加一个护理任务计划")

    hospitalization = await session.get(Hospitalization, hosp_id)
    if not hospitalization:
        raise HTTPException(status_code=404, detail="住院记录不存在")
    if hospitalization.hosp_doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能为自己负责的住院患者添加护理任务")
    if hospitalization.status != "在院":
        raise HTTPException(status_code=400, detail="仅能为在院患者添加护理任务")
    if not hospitalization.ward_id:
        raise HTTPException(status_code=400, detail="该住院记录缺少病房信息")

    now = datetime.now()
    created_tasks: List[NurseTask] = []
    for plan in payload.plans:
        if plan.start_time <= now:
            raise HTTPException(status_code=400, detail=f"{plan.type} 的开始时间需晚于当前时间")

        schedule_times = expand_task_schedule(plan)
        if not schedule_times:
            raise HTTPException(status_code=400, detail=f"{plan.type} 未生成任何护理任务，请检查频次设置")

        snapshot = json.dumps([item.dict() for item in plan.medicines], ensure_ascii=False) if plan.medicines else None
        detail_text = build_task_detail(plan)

        for scheduled_time in schedule_times:
            if scheduled_time <= now:
                raise HTTPException(status_code=400, detail=f"{plan.type} 的执行时间需晚于当前时间")
            service_fee = estimate_service_fee(plan.type)
            task = NurseTask(
                type=plan.type,
                time=scheduled_time,
                hosp_id=hosp_id,
                detail=detail_text,
                medicine_snapshot=snapshot,
                service_fee=service_fee,
            )
            session.add(task)
            created_tasks.append(task)

    if not created_tasks:
        raise HTTPException(status_code=400, detail="未生成任何护理任务，请检查任务计划配置")

    await session.commit()
    for task in created_tasks:
        await session.refresh(task)
    return {"created": len(created_tasks), "tasks": created_tasks}

# --- 新：根据挂号ID查询对应病历（仅限接诊该挂号的医生） ---
@router.get("/consultations/{reg_id}/record", response_model=MedicalRecord)
async def get_record_by_reg(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能查看自己的病人")

    rec_stmt = select(MedicalRecord).where(MedicalRecord.reg_id == reg_id)
    record = (await session.execute(rec_stmt)).scalars().first()
    if not record:
        raise HTTPException(status_code=404, detail="病历不存在")
    return record


# 新：根据挂号ID查询挂号信息与患者信息（仅限接诊该挂号的医生）
@router.get("/consultations/{reg_id}/info")
async def get_consultation_info(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能查看自己的挂号信息")

    patient = await session.get(Patient, registration.patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="关联的患者不存在")

    return {"registration": registration, "patient": patient}


@router.get("/dept/doctors")
async def list_my_department_doctors(
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = (
        select(Doctor)
        .join(UserAccount, UserAccount.phone == Doctor.phone)
        .where(Doctor.dept_id == doctor.dept_id)
        .where(UserAccount.status == "启用")
    )
    doctors = (await session.execute(stmt)).scalars().all()
    return [
        {
            "doctor_id": d.doctor_id,
            "name": d.name,
            "title": d.title,
        }
        for d in doctors
    ]


@router.post("/consultations/{reg_id}/hospitalize")
async def hospitalize_patient(
    reg_id: int,
    payload: HospitalizePayload,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能为自己的病人办理住院")

    rec_stmt = select(MedicalRecord).where(MedicalRecord.reg_id == reg_id)
    record = (await session.execute(rec_stmt)).scalars().first()
    if not record:
        raise HTTPException(status_code=404, detail="请先建立病历再办理住院")

    # 严格保证单患者仅有一条“在院”住院单
    active_stmt = (
        select(Hospitalization, MedicalRecord, Registration)
        .join(MedicalRecord, Hospitalization.record_id == MedicalRecord.record_id)
        .join(Registration, MedicalRecord.reg_id == Registration.reg_id)
        .where(Registration.patient_id == registration.patient_id)
        .where(Hospitalization.status == "在院")
    )
    active_row = (await session.execute(active_stmt)).first()
    if active_row:
        active_hosp, active_record, active_reg = active_row
        if active_reg.reg_id != registration.reg_id:
            raise HTTPException(status_code=400, detail="该患者已有在院住院单，无法重复办理")
        # 同一挂号更换病房：若尚未完成办理，删除旧住院单；否则标记出院再新建
        if registration.status != RegStatus.FINISHED:
            await session.delete(active_hosp)
        else:
            active_hosp.status = "已出院"
            active_hosp.out_date = datetime.now()
        await session.flush()

    ward = await session.get(Ward, payload.ward_id)
    if not ward or ward.dept_id != doctor.dept_id:
        raise HTTPException(status_code=400, detail="请选择当前科室的有效病房")

    count_stmt = select(func.count()).where(
        Hospitalization.ward_id == ward.ward_id,
        Hospitalization.status == "在院"
    )
    occupied = (await session.execute(count_stmt)).scalar_one() or 0
    if occupied >= ward.bed_count:
        raise HTTPException(status_code=400, detail="该病房已满，请选择其他病房")

    target_doc_id = payload.hosp_doctor_id or doctor.doctor_id
    target_doc_stmt = (
        select(Doctor)
        .join(UserAccount, UserAccount.phone == Doctor.phone)
        .where(Doctor.doctor_id == target_doc_id)
        .where(Doctor.dept_id == doctor.dept_id)
        .where(UserAccount.status == "启用")
    )
    target_doc = (await session.execute(target_doc_stmt)).scalars().first()
    if not target_doc:
        raise HTTPException(status_code=400, detail="请选择当前科室的有效住院医生")

    hosp = Hospitalization(
        ward_id=ward.ward_id,
        status="在院",
        hosp_doctor_id=target_doc.doctor_id,
        record_id=record.record_id
    )
    session.add(hosp)
    try:
        await session.commit()
    except IntegrityError as exc:
        await session.rollback()
        msg = str(getattr(exc, "orig", exc)).lower()
        if "uq_hosp_active_record_id" in msg or "duplicate entry" in msg:
            raise HTTPException(status_code=400, detail="该患者已在院，无法重复办理")
        raise HTTPException(status_code=400, detail="住院办理失败，请稍后再试")
    await session.refresh(hosp)
    return {"message": "住院办理完成", "hospitalization": hosp}


@router.get("/consultations/{reg_id}/transfer")
async def generate_transfer_form(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能为自己的病人申请转院")

    patient = await session.get(Patient, registration.patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="关联患者不存在")

    rec_stmt = select(MedicalRecord).where(MedicalRecord.reg_id == reg_id)
    record = (await session.execute(rec_stmt)).scalars().first()

    exams = []
    if record:
        exam_stmt = select(Examination).where(Examination.record_id == record.record_id)
        exams = (await session.execute(exam_stmt)).scalars().all()

    doc = Document()

    def add_text(para_text: str, bold: bool = False, size: int = 12, align: WD_ALIGN_PARAGRAPH = None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        run = p.add_run(para_text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    add_text("转院申请单", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_text(f"转入医院：____________________________")
    add_text(f"患者姓名：{patient.name or ''}")
    add_text(f"联系方式：{patient.phone or ''}")
    add_text(f"就诊日期：{registration.visit_date.strftime('%Y-%m-%d') if registration.visit_date else ''}")
    add_text("")

    add_text("【申请说明】", bold=True)
    add_text("因患者病情需要进一步诊治，特申请将患者转入上述医院继续治疗。请贵院协助接诊并提供后续诊疗支持。")
    add_text("")

    add_text("【病历摘要】", bold=True)
    if record:
        add_text(f"主诉：{record.complaint}")
        add_text(f"诊断：{record.diagnosis}")
        add_text(f"建议：{record.suggestion}")
    else:
        add_text("尚未填写病历")
    add_text("")

    add_text("【检查记录】", bold=True)
    if exams:
        for exam in exams:
            add_text(f"· {exam.type} · 结果：{exam.result} · {format_datetime(exam.date)}")
    else:
        add_text("暂未开具检查")

    add_text("")
    add_text("申请科室/医生签字：____________________________")
    add_text("日期：____________________________")

    payload = io.BytesIO()
    doc.save(payload)
    payload.seek(0)
    response = StreamingResponse(payload, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response.headers["content-disposition"] = f"attachment; filename=transfer_{reg_id}.docx"
    return response
# --- 新：开始办理（将 status 从 WAITING 改为 IN_PROGRESS） ---
@router.post("/consultations/{reg_id}/start")
async def start_handling(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能处理自己的病人")
    if registration.status != RegStatus.WAITING:
        raise HTTPException(status_code=400, detail="只有排队中的挂号可以开始办理")

    registration.status = RegStatus.IN_PROGRESS
    session.add(registration)
    await session.commit()
    await session.refresh(registration)
    return registration


# --- 新：完成办理（将 status 从 IN_PROGRESS 改为 FINISHED） ---
@router.post("/consultations/{reg_id}/finish")
async def finish_handling(
    reg_id: int,
    doctor: Doctor = Depends(get_current_doctor),
    session: AsyncSession = Depends(get_session)
):
    stmt = select(Registration).where(Registration.reg_id == reg_id)
    registration = (await session.execute(stmt)).scalars().first()
    if not registration:
        raise HTTPException(status_code=404, detail="挂号单不存在")
    if registration.doctor_id != doctor.doctor_id:
        raise HTTPException(status_code=403, detail="您只能处理自己的病人")
    if registration.status != RegStatus.IN_PROGRESS:
        raise HTTPException(status_code=400, detail="只有办理中的挂号可以完成办理")

    registration.status = RegStatus.FINISHED
    session.add(registration)
    await session.commit()
    await session.refresh(registration)

    return registration